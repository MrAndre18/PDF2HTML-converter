import re
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.shared import RGBColor
from docx.oxml.ns import qn

def convert_pdf_to_docx(pdf_path, docx_path):
  cv = Converter(pdf_path)
  cv.convert(docx_path, start=0, end=None)
  cv.close()
  
def remove_empty_paragraphs(doc):
  paragraphs = doc.paragraphs
  for paragraph in paragraphs:
    if not paragraph.text.strip():
      p = paragraph._element
      p.getparent().remove(p)

def remove_lines_with_text(doc, text):
  for paragraph in doc.paragraphs:
    if text in paragraph.text:
      paragraph.clear()

def remove_lines_with_datetime(doc):
  datetime_pattern = r'\d{2}\.\d{2}\.\d{4}, \d{2}:\d{2}'

  for paragraph in doc.paragraphs:
    if re.search(datetime_pattern, paragraph.text):
      paragraph.text = ""

def set_font_size(doc, size):
  for paragraph in doc.paragraphs:
    for run in paragraph.runs:
      run.font.size = Pt(size)

def set_font_name(doc, font_name):
  for paragraph in doc.paragraphs:
    for run in paragraph.runs:
      run.font.name = font_name
      
def set_font_color(doc, color):
  for paragraph in doc.paragraphs:
    for run in paragraph.runs:
      run.font.color.rgb = color  # Устанавливаем цвет текста

def convert_links_to_text(doc):
  for paragraph in doc.paragraphs:
    for run in paragraph.runs:
      hyperlink = run.element.find(qn("w:hyperlink"))
      if hyperlink is not None:
        run.clear()
        hyperlink_text = hyperlink.find(qn("w:r"))
        if hyperlink_text is not None:
          for element in hyperlink_text:
            if element.tag.endswith("t"):
              t = parse_xml(f'<w:t xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">{element.text}</w:t>')
              run._r.append(t)

pdf_path = 'PDF/Ministerial_Resolution_№_71_of_1989_Regarding_the_procedures_for.pdf'
docx_path = 'Result.docx'
text_to_remove = 'https://uae.saderlex.com/'

convert_pdf_to_docx(pdf_path, docx_path)
doc = Document(docx_path)

remove_lines_with_text(doc, text_to_remove)
remove_lines_with_datetime(doc)

set_font_size(doc, 14)
set_font_name(doc, 'Inter')
set_font_color(doc, RGBColor(0, 0, 0))
convert_links_to_text(doc)

remove_empty_paragraphs(doc)

doc.save('Result.docx')